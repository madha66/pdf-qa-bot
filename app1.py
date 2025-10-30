import os
import tempfile
import streamlit as st
import fitz
import hashlib
import base64
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.chains import RetrievalQA
from langchain_groq import ChatGroq
from docx import Document
from docx2pdf import convert


def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    return "".join(page.get_text() for page in doc)


def split_text(text, size=1000, overlap=100):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=size,
        chunk_overlap=overlap
    )
    return splitter.split_text(text)


def embed_text_with_faiss(chunks, index_path):
    embedder = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    vstore = FAISS.from_texts(chunks, embedding=embedder)
    vstore.save_local(index_path)
    return vstore


def get_llm_chain():
    groq_api_key = os.environ.get("GROQ_API_KEY")
    if not groq_api_key:
        raise ValueError("API Key not found")
    os.environ["GROQ_API_KEY"] = groq_api_key

    llm = ChatGroq(
        model_name="llama-3.1-8b-instant",
        temperature=0.3,
        max_tokens=300,
    )
    return llm


def get_answer(vstore, query, chain):
    retriever = vstore.as_retriever(search_kwargs={"k": 3})
    qa_bot = RetrievalQA.from_chain_type(
        llm=chain,
        retriever=retriever,
        chain_type="stuff",
        return_source_documents=False
    )
    result = qa_bot({"query": query})
    return result['result'].strip()


def create_docx(chat):
    doc = Document()
    doc.add_heading("Chatting with bot", 0)

    for q, a in chat:
        doc.add_paragraph(f"You: {q}", style='List Bullet')
        doc.add_paragraph(f"Bot: {a}", style="List Bullet")
        doc.add_paragraph(" ")

    docx_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(docx_file.name)
    return docx_file.name


def file_download(filepath, filename, label):
    with open(filepath, "rb") as f:
        bytes_data = f.read()
    encoded = base64.b64encode(bytes_data).decode()
    href = f'<a href="data:application/octet-stream;base64,{encoded}" download="{filename}">{label}</a>'
    return href


def get_file_hash(uploaded_file):
    return hashlib.md5(uploaded_file.getvalue()).hexdigest()


# ----------------------- Streamlit App -----------------------

st.set_page_config(page_title="PDF Q/A", layout="centered")
st.title("📚 Chat with your PDF")

if "chat" not in st.session_state:
    st.session_state.chat = []
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "chain" not in st.session_state:
    st.session_state.chain = None
if "last_file_hash" not in st.session_state:
    st.session_state.last_file_hash = ""

file = st.file_uploader("Upload a PDF", type="pdf")

if file:
    file_hash = get_file_hash(file)
    index_dir = os.path.join(tempfile.gettempdir(), f"faiss_{file_hash}")

    if st.session_state.last_file_hash != file_hash:
        with st.spinner("Processing PDF…"):
            st.session_state.chat = []

            if os.path.exists(index_dir):
                vstore = FAISS.load_local(
                    index_dir,
                    embeddings=HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2"),
                    allow_dangerous_deserialization=True
                )
            else:
                text = extract_text_from_pdf(file)

                if not text.split():
                    st.warning("The uploaded PDF contains no readable text. Please upload a readable PDF file.")
                    st.stop()
                else:
                    chunks = split_text(text)
                    if not chunks:
                        st.warning("The PDF was read, but no chunks could be created.")
                        st.stop()
                    else:
                        vstore = embed_text_with_faiss(chunks, index_dir)

            st.session_state.vectorstore = vstore
            st.session_state.chain = get_llm_chain()
            st.session_state.last_file_hash = file_hash
            st.success("Ready! Ask a question.")

if st.session_state.vectorstore and st.session_state.chain:
    with st.form("question-form"):
        if not file:
            q = st.text_input("Your question:", disabled=True)
            submit = st.form_submit_button("Ask", disabled=True)
        else:
            q = st.text_input("Your question:", disabled=False)
            submit = st.form_submit_button("Ask", disabled=False)

        if submit and q:
            with st.spinner("Thinking…"):
                a = get_answer(st.session_state.vectorstore, q, st.session_state.chain)
                if not a or a.lower() in ["i don't know", "i cannot answer that", ""]:
                    a = "Sorry, I couldn’t find the answer in the document."
                st.session_state.chat.append((q, a))

    with st.container():
        for q, a in st.session_state.chat:
            st.markdown(f"**You:** {q}")
            st.markdown(f"**Bot:** {a}")
            st.markdown("---")

    if st.session_state.chat:
        st.markdown("### Download Chat")
        docx = create_docx(st.session_state.chat)
        st.markdown(file_download(docx, "chat.docx", "Download as .docx"), unsafe_allow_html=True)
else:
    st.warning("Please upload a document to chat with the bot (only one document can be uploaded at a time).")
